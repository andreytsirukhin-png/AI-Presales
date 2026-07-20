"""Export proposal review reports."""

from io import BytesIO

from docx import Document

from app.modules.projects.schemas.review import ReviewReport, ReviewSeverity


def review_to_markdown(review: ReviewReport) -> str:
    """Render a review report as Markdown."""
    lines = [
        f"# Proposal Review — {review.project_name}",
        "",
        f"_Generated: {review.generated_at.isoformat()}_",
        "",
        "## Metrics",
        "",
        f"- Coverage: **{review.metrics.coverage_percent:.1f}%**",
        f"- Readiness score: **{review.metrics.readiness_score:.1f}**",
        f"- Requirements covered: {review.metrics.requirements_covered}",
        f"- Requirements missing: {review.metrics.requirements_missing}",
        f"- Critical findings: {review.metrics.critical_findings}",
        f"- Average severity: {review.metrics.average_severity_score:.2f}",
        "",
    ]

    for severity in ReviewSeverity:
        lines.append(f"## Findings — {severity.value.title()}")
        lines.append("")
        severity_findings = [
            finding
            for category in review.categories
            for finding in category.findings
            if finding.severity == severity
        ]
        if not severity_findings:
            lines.append("_None_")
            lines.append("")
            continue
        for finding in severity_findings:
            lines.append(f"### {finding.title}")
            lines.append("")
            lines.append(finding.description)
            lines.append("")
            lines.append(f"**Recommendation:** {finding.recommendation}")
            if finding.proposal_section:
                lines.append(f"**Proposal section:** {finding.proposal_section}")
            if finding.source_document:
                page = f" page {finding.page}" if finding.page else ""
                lines.append(f"**Source:** {finding.source_document}{page}")
            if finding.citations:
                lines.append("**Citations:**")
                for citation in finding.citations:
                    page = f" p.{citation.page}" if citation.page else ""
                    lines.append(f"- {citation.document}{page} (score {citation.score:.3f})")
            lines.append("")

    return "\n".join(lines).strip() + "\n"


def review_to_docx_bytes(review: ReviewReport) -> bytes:
    """Render a review report as DOCX bytes."""
    document = Document()
    document.add_heading(f"Proposal Review — {review.project_name}", level=0)
    document.add_paragraph(f"Coverage: {review.metrics.coverage_percent:.1f}%")
    document.add_paragraph(f"Readiness score: {review.metrics.readiness_score:.1f}")

    for category in review.categories:
        document.add_heading(category.title, level=1)
        document.add_paragraph(category.summary)
        for finding in category.findings:
            document.add_heading(finding.title, level=2)
            document.add_paragraph(finding.description)
            document.add_paragraph(f"Severity: {finding.severity.value}")
            document.add_paragraph(f"Recommendation: {finding.recommendation}")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
