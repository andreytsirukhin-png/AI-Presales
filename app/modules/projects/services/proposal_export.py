"""Export generated proposals to review formats."""

from io import BytesIO

from docx import Document

from app.modules.projects.schemas.proposal import Proposal


def proposal_to_markdown(proposal: Proposal) -> str:
    """Render a proposal as Markdown suitable for copy/export."""
    lines = [
        f"# Commercial Proposal — {proposal.project_name}",
        "",
        f"_Generated: {proposal.generated_at.isoformat()}_",
        "",
    ]
    for section in proposal.sections:
        lines.append(f"## {section.title}")
        lines.append("")
        lines.append(section.content.strip())
        lines.append("")
        if section.citations:
            lines.append("**Sources**")
            for citation in section.citations:
                page = f" page {citation.page}" if citation.page else ""
                lines.append(
                    f"- {citation.document}{page} (score {citation.score:.3f})"
                )
            lines.append("")
    return "\n".join(lines).strip() + "\n"


def proposal_to_docx_bytes(proposal: Proposal) -> bytes:
    """Render a proposal as a DOCX document."""
    document = Document()
    document.add_heading(f"Commercial Proposal — {proposal.project_name}", level=0)
    document.add_paragraph(f"Generated: {proposal.generated_at.isoformat()}")

    for section in proposal.sections:
        document.add_heading(section.title, level=1)
        document.add_paragraph(section.content.strip())
        if section.citations:
            document.add_paragraph("Sources:")
            for citation in section.citations:
                page = f" page {citation.page}" if citation.page else ""
                document.add_paragraph(
                    f"{citation.document}{page} (score {citation.score:.3f})",
                    style="List Bullet",
                )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def proposal_pdf_placeholder(proposal: Proposal) -> bytes:
    """Return a placeholder PDF export payload (Markdown bytes until PDF pipeline exists)."""
    note = (
        "PDF export placeholder.\n\n"
        + proposal_to_markdown(proposal)
    )
    return note.encode("utf-8")
